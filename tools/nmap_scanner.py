import argparse
import logging
import sys
import nmap
from docx import Document

def nmap_scan(target, verbosity):
    """Running nmap"""
    nm = nmap.PortScanner()
    verbosity_arg = f"-{'v' * verbosity}"
    logging.info("Scanning target: %s with verbosity level: %s", target, verbosity_arg)
    results = []
    try:
        nm.scan(target, arguments=f'-sC -sV {verbosity_arg}')
        for host in nm.all_hosts():
            result = f"Results for {host}:\n"
            for proto in nm[host].all_protocols():
                result += f"Protocol: {proto}\n"
                ports = nm[host][proto].keys()
                for port in ports:
                    port_info = nm[host][proto][port]
                    product = port_info.get('product', 'Unknown')
                    version = port_info.get('version', '')
                    result += f"  Port {port}: {port_info['state']} - {product} {version}\n"
            results.append(result)
    except nmap.PortScannerError as e:
        logging.error("An error occurred during scanning: %s", e)
        sys.exit(1)
    return results

def setup_logging():
    """setup logging"""
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def parse_args():
    """nmap scanning script"""    
    parser = argparse.ArgumentParser(description='Nmap Scanning Script with Verbosity Levels')
    parser.add_argument('target', type=str, help='Target to scan (IP address or network range)')
    parser.add_argument('-v', '--verbosity',
    action='count',
    default=0,
    help='Increase output verbosity')

    args = parser.parse_args()

    if args.verbosity > 3:
        logging.warning("Maximum verbosity level is 3 (-vvv), using maximum verbosity.")
        args.verbosity = 3

    return args

def create_docx_report(results):
    """Create a Word document report"""
    doc = Document()
    doc.add_heading('Nmap Scan Report', level=1)
    for result in results:
        doc.add_paragraph(result)
        doc.add_page_break()
    doc.save('nmap_scan_report.docx')

if __name__ == "__main__":
    args = parse_args()
    setup_logging()
    scan_results = nmap_scan(args.target, args.verbosity)
    create_docx_report(scan_results)
    logging.info("Report generated: nmap_scan_report.docx")
